"""
document_export.py – Esporta un file DOCX strutturato dai risultati della pipeline video indexer.
"""

from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches
import logging

from core.config import OLLAMA_LLM_MODEL, OLLAMA_URL, FRAMES_DIR

from utils.ollama import generate_title_and_abstract

log = logging.getLogger("document_export")

def export_docx(
    indicizzazione: dict,
    output_path: Path,
    title: Optional[str] = None,
    abstract: Optional[str] = None
):
    """
    Crea un file DOCX con titolo, abstract, segmenti con testo e immagini + didascalia.
    - indicizzazione: dict prodotto dalla pipeline (video, segments)
    - output_path: dove salvare il file DOCX
    - title, abstract: opzionali, se vuoti vengono generati con Ollama
    """
    doc = Document()
    segments = indicizzazione["segments"]
    video_name = indicizzazione.get("video", "Video")
    images_dir = FRAMES_DIR / Path(video_name).stem

    # Generazione titolo e abstract via Ollama se mancano
    if not title or not abstract:
        title_gen, abstract_gen = generate_title_and_abstract(
            segments
        )
        title = title or title_gen
        abstract = abstract or abstract_gen

    # Titolo e abstract
    doc.add_heading(title, 0)
    if abstract:
        doc.add_paragraph(abstract)
        doc.add_paragraph("")  # Spazio

    # Segmenti temporali ordinati
    for idx, segment in enumerate(segments):
        interv = segment.get("interval", {})
        seg_title = f"Segmento {idx+1} ({interv.get('start', '?')}–{interv.get('end', '?')} s)"
        doc.add_heading(seg_title, level=1)
        doc.add_paragraph(segment.get("fulltext", ""))

        # Screenshot con didascalie
        for sc in segment.get("screenshots", []):
            image_name = sc.get("image")
            image_path = images_dir / image_name
            descr = sc.get("description", "")
            if image_path.exists():
                try:
                    doc.add_picture(str(image_path), width=Inches(4.5))  # Adatta larghezza a piacere
                except Exception as e:
                    log.warning(f"Non riesco a inserire l'immagine {image_path}: {e}")
            if descr:
                doc.add_paragraph(descr, style="Caption")
        doc.add_paragraph("")  # Spazio tra segmenti

    doc.save(str(output_path))
    log.info(f"Documento DOCX esportato in {output_path}")

# Esempio d'uso
if __name__ == "__main__":
    import json

    indicizzazione_json = Path("indicizzazione_output.json")
    output_docx = Path("report_video.docx")

    with open(indicizzazione_json, "r", encoding="utf-8") as f:
        indicizzazione = json.load(f)

    export_docx(
        indicizzazione=indicizzazione,
        output_path=output_docx,
        title=None,
        abstract=None
    )
